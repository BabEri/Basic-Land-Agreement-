from docx import Document
import re
from datetime import date
from random import choice
from docx.enum.text import WD_ALIGN_PARAGRAPH,  WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from currency_notes import matcher

iu = Document()
style = iu.styles
title_style = style.add_style('ETitle', WD_STYLE_TYPE.CHARACTER)
title_font = title_style.font
title_font.size = Pt(16)
title_font.name = 'Times New Roman'
main_style = style.add_style('EMain', WD_STYLE_TYPE.CHARACTER)
main_font = main_style.font
main_font.size = Pt(14)
main_font.name = 'Garamond'
mainHead_style = style.add_style('mainHead', WD_STYLE_TYPE.CHARACTER)
mainHead_font = mainHead_style.font
mainHead_font.size = Pt(15)
mainHead_font.name = 'Garamond'
jurat_style = style.add_style('EJurat', WD_STYLE_TYPE.CHARACTER)
jurat_font = jurat_style.font
jurat_font.size = Pt(11)
jurat_font.name = 'Cambria'

#t = date.today() or t = date(Y,M,D) for specific dates
t = date.today()
def position(a):
	if (a == 1 or (a%10 == 1)) and str(a)[-2:] != '11':
	 	a = f'{a}st'
	elif (a== 2 or (a%10)==2) and str(a)[-2:] != '12':
	 	a = f'{a}nd'
	elif (a == 3 or (a%10)==3) and str(a)[-2:] != '13' :
	 	a = f'{a}rd'
	else:
		a = f'{a}th'
	return a
def today(k):#this is to rectify the problem in converting python datetime moodule to appropriately state '15th day of October, 1995'
	a = k.strftime(f'%d' + position(k.day)[-2:])
	if k.day < 10:
		a = f'{a[1]}' + position(k.day)[-2:]
	return a + k.strftime(f' of %B, %Y')
part = choice([('one', 'other'), ('first','second')])#to make sure the agreement has a bit of flexibility
s = input ('name of solicitor: ')
solicitor = s.title() + ' Esq.'
vendor_list = []
no_of_vendor = int (input('number of vendors: '))
address_of_vendor = []
display_vendor = 'Vendor'
if no_of_vendor > 1:
	for i in range(no_of_vendor):
		a = (input(f'enter the name of the {position(i + 1)} Vendor: ')).title()
		vendor_list.append(a)
		b = input(f'enter the address of the {position(i + 1)} Vendor: ')
		address_of_vendor.append(b)
		display_vendor = 'Vendors'
else:
	a = (input(f'enter the name of the Vendor: ')).title()
	vendor_list.append(a)
	b = input(f'enter the address of the Vendor: ')
	address_of_vendor.append(b)
		
address_of_vendor = [i.title() for i in address_of_vendor if len(i) > 3]
vendor_dict = dict(zip(vendor_list,address_of_vendor)) 
buyer_list = []
no_of_buyer = int (input('number of buyers: '))
address_of_buyer = []
display_buyer = 'Puchaser'
if no_of_buyer > 1:
	for i in range(no_of_buyer):
		a = (input(f'enter the name of the {position(i + 1)} buyer: ')).title()
		buyer_list.append(a)
		b = input(f'enter the address of the {position(i + 1)} buyer: ')
		address_of_buyer.append(b)
		display_buyer = 'Puchasers'
else:
		a = (input(f'enter the name of the Buyer: ')).title()
		buyer_list.append(a)
		b = input(f'enter the address of the  Buyer: ')
		address_of_buyer.append(b)

buyer_dict = dict(zip(buyer_list,address_of_buyer))
if no_of_buyer > 1 and no_of_vendor > 2 :
	recital = (f'''Now the Agreement recites as follows:
I. The {display_vendor} are beneficial owners of the property 
II. The {display_vendor} are desirous of alienating the property and the {display_buyer} are also desirous of buying the property.
III.The Parties hereby make this agreement to govern thier relationship.
''',
f'''RECITAL
I. The {display_vendor} are the beneficial owners of the property.
II. The property is a family property.
III. {vendor_list[0]} is the Head of the family while {' and '.join(vendor_list[1:]) } are principal members
IV. The {display_vendor} are desirous of alienating the property and the {display_buyer} are also desirous of buying the property.
V. The Parties hereby make this agreement to govern thier relationship.
''')
elif no_of_buyer > 1 and no_of_vendor == 2 :
	recital = (f'''Now the Agreement recites as follows:
I. The {display_vendor} are beneficial owners of the property 
II. The {display_vendor} are desirous of alienating the property and the {display_buyer} are also desirous of buying the property.
III.The Parties hereby make this agreement to govern thier relationship.
''',
f'''RECITAL
I. The {display_vendor} are the beneficial owners of the property.
II. The property is a family property.
III. {vendor_list[0]} is the Head of the family while {' and '.join(vendor_list[1:]) } is a principal member of their family
IV. The {display_vendor} are desirous of alienating the property and the {display_buyer} are also desirous of buying the property.
V. The Parties hereby make this agreement to govern thier relationship.
''')
elif no_of_buyer == 1 and no_of_vendor == 2 :
	recital = (f'''Now the Agreement recites as follows:
I. The {display_vendor} are beneficial owners of the property 
II. The {display_vendor} are desirous of alienating the property and the {display_buyer} is also desirous of buying the property.
III.The Parties hereby make this agreement to govern thier relationship.
''',
f'''RECITAL
I. The {display_vendor} are the beneficial owners of the property.
II. The property is a family property.
III. {vendor_list[0]} is the Head of the family while {' and '.join(vendor_list[1:]) } is a principal member of their family
IV. The {display_vendor} are desirous of alienating the property and the {display_buyer} is also desirous of buying the property.
V. The Parties hereby make this agreement to govern thier relationship.
''')
elif no_of_buyer > 1 and no_of_vendor == 1 :
	recital = (f'''Now the Agreement recites as follows:
I. The {display_vendor} is the beneficial owner of the property 
II. The {display_vendor} is desirous of alienating the property and the {display_buyer} are also desirous of buying the property.
III.The Parties hereby make this agreement to govern thier relationship.
''',
f'''RECITAL
I. The {display_vendor} is the beneficial owner of the property.
II. The property is a family property.
III. {vendor_list[0]} is the Head of the family who has the authority of the principal members to alienate
IV. The {display_vendor} on behalf of the family is desirous of alienating the property and the {display_buyer} are also desirous of buying the property.
V. The Parties hereby make this agreement to govern thier relationship.
''')
elif no_of_buyer == 1 and no_of_vendor == 1 :
	recital = (f'''Now the Agreement recites as follows:
I. The {display_vendor} is the beneficial owner of the property 
II. The {display_vendor} is desirous of alienating the property and the {display_buyer} is also desirous of buying the property.
III.The Parties hereby make this agreement to govern thier relationship.
''',
f'''RECITAL
I. The {display_vendor} is the beneficial owner of the property.
II. The property is a family property.
III. {vendor_list[0]} is the Head of the family who has the authority of the principal members to alienate
IV. The {display_vendor} on behalf of the family is desirous of alienating the property and the {display_buyer} is also desirous of buying the property.
V. The Parties hereby make this agreement to govern thier relationship.
''')
elif no_of_buyer == 1 and no_of_vendor > 1 :
	recital = (f'''Now the Agreement recites as follows:
I. The {display_vendor} are beneficial owners of the property 
II. The {display_vendor} are desirous of alienating the property and the {display_buyer} is also desirous of buying the property.
III.The Parties hereby make this agreement to govern thier relationship.
''',
f'''RECITAL
I. The {display_vendor} are the beneficial owners of the property.
II. The property is a family property.
III. {vendor_list[0]} is the Head of the family while {' and '.join(vendor_list[1:]) } are principal members
IV. The {display_vendor} are desirous of alienating the property and the {display_buyer} is also desirous of buying the property.
V. The Parties hereby make this agreement to govern thier relationship.
''')

description_of_property = input ('describe the property: ')
consideration = input('how much is the property: ')
type_of_property = (input('is property a family property, input yes or no: ')).lower()
is_illiterate = input('is any of the parties an illiterate, input yes or no: ') 
illiterate_jurat = ''' 
The content of this agreement was read and expalained by 
____________________________ in YORUBA Language 
to the understanding of 
'''
def pack(**kwargs):
	s = ''
	for k in kwargs:
		s+= f'{k} \n'
	return s

def unpack(**kwargs):
	s = ''
	for k in kwargs:
		s += f'{k} of {kwargs[k]} and '
	s = s[:-4]
	return s
i = 0 #for numbering
def num():
	global i
	i += 1
	return i

def deposit():
	money_deposit = input('enter the amount of deposit: ')
	mw = '' 
	c = deposit_paragraph = iu.add_paragraph()
	a = deposit_paragraph.add_run(f'{num()}. DEPOSIT', style = 'mainHead').bold = True
	d = deposit_paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
	b = deposit_paragraph_1 = iu.add_paragraph()
	if no_of_buyer > 1:
		e = deposit_paragraph_1.add_run(
f'''a. The {display_buyer.capitalize()} undertake to deposit {money_deposit} {matcher(money_deposit)} to {solicitor}\t
b. The {display_buyer.capitalize()} further agree to forfeit the deposit if they fails to pay the consideration in full and Paragraph 1 of this Agreement shall apply
		''', style = 'EMain')
	e = deposit_paragraph_1.add_run(
f'''a. The {display_buyer.capitalize()} undertakes to deposit {money_deposit} {matcher(money_deposit)} to {solicitor} \t
b. The {display_buyer.capitalize()} further agrees to forfeit the deposit if s/he fails to pay the consideration in full and Paragraph 1 of this Agreement shall apply\t
		''', style = 'EMain')
	f = deposit_paragraph_1.paragraph_format.alignment =WD_ALIGN_PARAGRAPH.JUSTIFY
	mw += f'{c}'
	mw += f'{a}'
	mw += f'{d}'
	mw += f'{b}'
	mw += f'{e}'
	mw += f'{f}'
	return mw
def execution(tx,fg):
	execution_clause = ''
	a = execution_paragraph = iu.add_paragraph()
	b = execution_paragraph.add_run(f'''\n
The Parties hereby execute this agreement in the manner below and on the date stated above \t
 ''', style = 'EMain')
	c = execution_paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
	execution_clause += f'{a}'
	execution_clause += f'{b}'
	execution_clause += f'{c}'
	if len (tx) > 1:
		for i in range(len(tx)):
			d= execution_paragraph_1 = iu.add_paragraph()
			execution_clause += f'{d}'
			e = execution_paragraph_1.add_run(f'SIGNED BY THE {position(i+1)} VENDOR', style = 'EMain')
			execution_clause += f'{e}'
			f = execution_paragraph_2 = iu.add_paragraph()
			execution_clause += f'{f}'
			g = execution_paragraph_2.add_run('_'* len(tx[i]), style = 'EMain')
			execution_clause = f'{g}'
			h = execution_paragraph_2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
			execution_clause += f'{h}'
			j = execution_paragraph_3 = iu.add_paragraph()
			execution_clause += f'{j}'
			k = execution_paragraph_3.add_run(f'{tx[i]}',style = 'EMain')
			execution_clause += f'{k}'
			l = execution_paragraph_3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
			execution_clause += f'{l}'
	else:
		d= execution_paragraph_1 = iu.add_paragraph()
		execution_clause += f'{d}'
		e = execution_paragraph_1.add_run(f'SIGNED BY THE VENDOR', style = 'EMain')
		execution_clause += f'{e}'
		f = execution_paragraph_2 = iu.add_paragraph()
		execution_clause += f'{f}'
		g = execution_paragraph_2.add_run('_'* len(tx[0]), style = 'EMain')
		execution_clause = f'{g}'
		h = execution_paragraph_2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
		execution_clause += f'{h}'
		j = execution_paragraph_3 = iu.add_paragraph()
		execution_clause += f'{j}'
		k = execution_paragraph_3.add_run(f'{tx[0]}',style = 'EMain') 
		execution_clause += f'{k}'
		l = execution_paragraph_3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
		execution_clause += f'{l}'																			
	if len(fg) > 1:
		for  i in range (len(fg)):
			m = execution_paragraph_1 = iu.add_paragraph()
			execution_clause += f'{m}'
			n = execution_paragraph_1.add_run(f'SIGNED BY THE {position(i+1)} BUYER', style = 'EMain')
			execution_clause += f'{n}'
			o = execution_paragraph_2 = iu.add_paragraph()
			execution_clause += f'{o}'
			p = execution_paragraph_2.add_run('_'* len(fg[i]), style = 'EMain')
			execution_clause = f'{p}'
			q = execution_paragraph_2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
			execution_clause += f'{q}'
			r = execution_paragraph_3 = iu.add_paragraph()
			execution_clause += f'{r}'
			s = execution_paragraph_3.add_run(f'{fg[i]}',style = 'EMain')
			execution_clause += f'{s}'
			t = execution_paragraph_3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
			execution_clause += f'{t}'
	else:
		m = execution_paragraph_1 = iu.add_paragraph()
		execution_clause += f'{m}'
		n = execution_paragraph_1.add_run(f'SIGNED BY THE BUYER', style = 'EMain')
		execution_clause += f'{n}'
		o = execution_paragraph_2 = iu.add_paragraph()
		execution_clause += f'{o}'
		p = execution_paragraph_2.add_run('_'* len(fg[0]), style = 'EMain')
		execution_clause = f'{p}'
		q = execution_paragraph_2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
		execution_clause += f'{q}'
		r = execution_paragraph_3 = iu.add_paragraph()
		execution_clause += f'{r}'
		s = execution_paragraph_3.add_run(f'{fg[0]}',style = 'EMain')
		execution_clause += f'{s}'
		t = execution_paragraph_3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
		execution_clause += f'{t}'
	return execution_clause
name_of_witness = []
job_of_witness = []
address_of_witness =  []
for i in range (2):
	a = (input(f'enter the name of the {position(i + 1)} Witness: ')).title()
	name_of_witness.append(a)
	b = (input(f'enter the address of the {position(i + 1)} Witness: ')).title()
	address_of_witness.append(b)
	c = (input(f'enter the occupation of the {position(i + 1)} Witness: ')).title()
	job_of_witness.append(c)

docx_name = f'Landsale Agreement between {vendor_list[0]} and {buyer_list[0]} {today(t)}.docx'
def make_agreement(recite= None, depos = False):
	with open( docx_name, 'w') as f:
		title_paragraph = iu.add_paragraph('\n\n\n')
		title_paragraph.add_run( f'''LAND SALE AGREEMENT\n\n BETWEEN \n\n {pack(**vendor_dict)} ({(display_vendor.upper())})\n\n AND
\n {pack(**buyer_dict)} ({(display_buyer.upper())})\n\n\n\n ''', style = 'ETitle').bold = True
		title_paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		solicitor_paragraph = iu.add_paragraph()
		solicitor_paragraph.add_run (f'PREPARED BY {solicitor} \n \n \n \n ', style = 'ETitle').bold = True
		solicitor_paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
		iu.add_page_break()
		introduction_paragraph = iu.add_paragraph()
		introduction_paragraph.add_run(f'''This Landsale Agreement is made this {today(t)} between {unpack(**vendor_dict)} (hereinafter referred to as {display_vendor}) of the {part[0]} part\n AND\n{unpack(**buyer_dict)} (hereinafter referred to as {display_buyer}) of the {part[1]} part \t''', style = 'EMain')
		introduction_paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
		if recite == 'family':
 			recit = iu.add_paragraph()
 			recit.add_run(f'{recital[1]}', style= 'EMain')
		elif recite:
 			recit = iu.add_paragraph()
 			recit.add_run(f'{recital[0]}', style= 'EMain')
		main1 = iu.add_paragraph()
		main1.add_run(f'Now the Agreement ', style='EMain')
		main1.add_run(f' WITNESSES ', style='EMain').bold = True
		main1.add_run(f'as follows:\t\n', style='EMain')
		if no_of_vendor > 1:
			main1.add_run (
f'''The {display_vendor} hereby agree to SELL all that {description_of_property}(hereinafter referred to as 'the property') to the {display_buyer} for a consideration of {consideration} ({matcher(consideration)}) in the following terms: \t''', style= 'EMain')
		main1.add_run (
f'''The {display_vendor} hereby agrees to SELL all that {description_of_property}(hereinafter referred to as 'the property') to the {display_buyer} for a consideration of {consideration} ({matcher(consideration)}) in the following terms: \t''', style= 'EMain')
		main1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
		main2 = iu.add_paragraph() 
		main2.add_run(f'{num()}. PAYMENT OF CONSIDERATION', style ='mainHead').bold = True
		main2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
		main21 = iu.add_paragraph()
		if no_of_buyer > 1 and no_of_vendor > 1:
			main21.add_run (
f'''a.The {display_buyer.capitalize()} agree to pay the consideration in full at the execution of this agreement or at least before they take possession of the property.\t
b. In the default of payment of the full consideration within 3 months of the execution of this agreement, the {display_buyer.capitalize()} shall be in breach of this contract and the {display_vendor.capitalize()} can either repudiate the contract or institute a case in the appropriate Court for breach of contract.\t
c. in the event that the {display_vendor.capitalize()} repudiate the contract, the {display_vendor.capitalize()} shall return any consideration paid by the {display_buyer.capitalize()} save any 'deposit\t'
''', style ='EMain' )
		elif no_of_buyer == 1 and no_of_vendor > 1:
			main21.add_run (
f'''a.The {display_buyer.capitalize()} agrees to pay the consideration in full at the execution of this agreement or at least before they take possession of the property.\t
b. In the default of payment of the full consideration within 3 months of the execution of this agreement, the {display_buyer.capitalize()} shall be in breach of this contract and the {display_vendor.capitalize()} can either repudiate the contract or institute a case in the appropriate Court for breach of contract.\t
c. in the event that the {display_vendor.capitalize()} repudiate the contract, the {display_vendor.capitalize()} shall return any consideration paid by the {display_buyer.capitalize()} save any 'deposit'\t
''', style ='EMain' )
		elif no_of_buyer > 1 and no_of_vendor == 1:
			main21.add_run (
f'''a.The {display_buyer.capitalize()} agree to pay the consideration in full at the execution of this agreement or at least before they take possession of the property.
b. In the default of payment of the full consideration within 3 months of the execution of this agreement, the {display_buyer.capitalize()} shall be in breach of this contract and the {display_vendor.capitalize()} can either repudiate the contract or institute a case in the appropriate Court for breach of contract.
c. in the event that the {display_vendor.capitalize()} repudiates the contract, the {display_vendor.capitalize()} shall return any consideration paid by the {display_buyer.capitalize()} save any 'deposit'
''', style ='EMain' )
		main21.add_run (
f'''a.The {display_buyer.capitalize()} agrees to pay the consideration in full at the execution of this agreement or at least before they take possession of the property.\t
b. In the default of payment of the full consideration within 3 months of the execution of this agreement, the {display_buyer.capitalize()} shall be in breach of this contract and the {display_vendor.capitalize()} can either repudiate the contract or institute a case in the appropriate Court for breach of contract.\t
c. in the event that the {display_vendor.capitalize()} repudiates the contract, the {display_vendor.capitalize()} shall return any consideration paid by the {display_buyer.capitalize()} save any 'deposit'\t
''', style ='EMain' )
		main21.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
		if depos:
			deposit()
		main3 = iu.add_paragraph()
		main3.add_run(f'{num()}. INDEMNIFICATION', style = 'mainHead').bold = True
		main3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
		main31 = iu.add_paragraph()
		if no_of_vendor == 1:
			main31.add_run(f'The {display_vendor} undertakes that the Property is free from any encumberance and will indemnify the {display_buyer} from any loss arising from any third party claim.', style = 'EMain') 
		main31.add_run(f'The {display_vendor} undertake that the Property is free from any encumberance and will indemnify the {display_buyer} from any loss arising from any third party claim.', style = 'EMain') 
		main31.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
		main4 = iu.add_paragraph()
		main4.add_run(f'{num()}. PERFECTION', style ='mainHead').bold = True
		main41 =  iu.add_paragraph()
		if no_of_vendor == 1:
			main41.add_run(f'The {display_vendor.capitalize()} undertakes to execute any document to perfect the transfer of the property to the {display_buyer}.', style = 'EMain')
		main41.add_run(f'The {display_vendor.capitalize()} undertake to execute any document to perfect the transfer of the property to the {display_buyer}.', style = 'EMain')
		main41.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
		main5 = iu.add_paragraph()
		main5.add_run(f'{num()}. APPLICABLE LAW', style = 'mainHead').bold = True
		main5.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
		main51 = iu.add_paragraph()
		main51.add_run(
 f'The Parties agree that this agreement shall be governed by the applicable Nigerian law, any dispute, controversies or ambiguity shall be interpreted by a Court of competent jurisdiction.\t', style = 'EMain')
		main51.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
		if is_illiterate == 'yes':
			illi = iu.add_paragraph()
			illi.add_run(illiterate_jurat, style = 'EMain').italic = True
			illi.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
		execution(vendor_list, buyer_list)
		witness_paragraph = iu.add_paragraph()
		witness_paragraph.add_run('IN THE PRESENCE OF :').bold = True
		witness_paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
		for i in range (2):
			witness_paragraph_1 = iu.add_paragraph()
			witness_paragraph_1.add_run(f'Name:       {name_of_witness[i]} \t').bold = True
			witness_paragraph_1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
			witness_paragraph_2 = iu.add_paragraph()
			witness_paragraph_2.add_run(f'address:       {address_of_witness[i]} \t').bold = True
			witness_paragraph_2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
			witness_paragraph_3 = iu.add_paragraph()
			witness_paragraph_3.add_run(f'Occupation:       {job_of_witness[i]} \t').bold = True
			witness_paragraph_3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
		iu.save(docx_name)
if type_of_property == 'yes':
	make_agreement(recite = 'family', depos = False)
else:
	make_agreement(recite = '', depos = False)







